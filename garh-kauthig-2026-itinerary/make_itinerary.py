#!/usr/bin/env python3
"""
Programme / running order for Garh Kauthig 2026.

Builds two outputs from the single data block below, so they cannot drift apart:

    garh-kauthig-2026-itinerary.docx   editable in Word / Google Docs / Pages
    garh-kauthig-2026-itinerary.md     editable as plain text
    garh-kauthig-2026-itinerary.html   designed A4 sheet -> PDF via render.mjs

EDIT THE DATA BLOCK, then run:  python3 make_itinerary.py

Everything here is transcribed from the organisers' notes. Where the notes were
ambiguous the item is left marked "to confirm" rather than guessed at, and the
open questions are listed in QUERIES so they appear on the sheet itself.

No timings have been invented. The poster gives a 01:00 PM start and nothing
further, so the Time column is deliberately blank for the organisers to fill.
"""

import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# =============================================================================
#  DATA  --  edit freely
# =============================================================================

EVENT = {
    'name': 'Garh Kauthig 2026',
    'theme': 'Tradition, Togetherness, and Timeless Culture',
    'date': 'Thursday, 13 August 2026',
    'time': '01:00 PM onwards',
    'venue': 'Silver Jubilee Convention Centre',
    'venue_sub': 'Graphic Era (Deemed to be University), Dehradun, Uttarakhand',
    'programme': 'Held as part of the EMERGE Induction Program 2026',
    'hosts': 'Graphic Era School of Management  ·  Swaragini, The Cultural Society '
             'of Graphic Era University',
}

TBC = '\u2014'          # an em dash: nothing invented, fill this in

# (Time, Item, Details, Presented by)
RUNNING_ORDER = [
    ('01:00 PM', 'Deep Prajwalan',
     'Ceremonial lighting of the lamp to open the programme',
     'Guests and Faculty'),

    (TBC, 'Devbhoomi', 'To confirm: item title, or the performing team', TBC),

    (TBC, 'Introducing Uttarayini', 'Introduction of the group before the performances',
     'Anchors'),

    (TBC, 'MBA Performance \u2014 First Year', '', 'MBA, First Year'),
    (TBC, 'MBA Performance \u2014 Second Year', '', 'MBA, Second Year'),

    (TBC, 'Performance 1', 'To confirm: item and performers', TBC),
    (TBC, 'Performance 2', 'To confirm: item and performers', TBC),
    (TBC, 'Performance 3', 'To confirm: item and performers', TBC),

    (TBC, 'Nandas', 'To confirm: item title, or the performing team', TBC),
    (TBC, 'Ramailo Cha', 'To confirm: item title, or the performing team', TBC),

    (TBC, 'Performance 4', 'To confirm: item and performers', TBC),
    (TBC, 'Performance 5', 'To confirm: item and performers', TBC),
    (TBC, 'Performance 6', 'To confirm: item and performers', TBC),

    (TBC, 'Uttarayini Performance', 'Concluding performance', 'Uttarayini'),

    (TBC, 'Awards and Prize Distribution', 'See the awards list below',
     'The Registrar'),

    (TBC, 'Conclusion', 'Close of the programme', TBC),
]

GUESTS = [
    ('Mr. Anshul Negi', 'Guest \u2014 Judge'),
    ('Miss Anjali', 'Guest \u2014 Judge'),
    ('Shivam', 'Guest \u2014 Judge'),
    ('Uday', 'Guest \u2014 not judging'),
]

AWARDS_BY = 'Presented by the Registrar'
AWARDS = [
    'First Prize',
    'Second Prize',
    'Third Prize',
    'Best Performance',
    'Certificate of Appreciation \u2014 to the team',
]

QUERIES = [
    'Timings: the poster gives only the 01:00 PM start, so no running times have '
    'been assumed. Tell me the finish time or the length of each slot and I will '
    'fill the whole column in.',
    'Performances 1 to 6 are unnamed in the notes \u2014 please add the item and the '
    'performers for each.',
    'Devbhoomi, Nandas and Ramailo Cha: are these item titles or the names of the '
    'performing teams?',
    '"Uttarayini (sth mein introduce)" has been read as: Uttarayini is introduced '
    'before the performances begin, and performs at the end. Please confirm the '
    'placement of the introduction.',
    '"Deep Prajwalan by guests and facilities" has been read as guests and '
    '<b>faculty</b>.',
    'Anshul Negi, Anjali and Shivam are shown as judges because Uday is marked '
    '"not judge" \u2014 please confirm, and add titles or surnames for Shivam and Uday.',
    'The anchors / compères are not named in the notes.',
]

# =============================================================================

HERE = os.path.dirname(os.path.abspath(__file__))
MAROON = RGBColor(0x6F, 0x11, 0x19)
CRIMSON = RGBColor(0x9C, 0x1C, 0x22)
COPPER = RGBColor(0xA3, 0x5A, 0x1D)
INK = RGBColor(0x33, 0x24, 0x1A)
GREY = RGBColor(0x60, 0x60, 0x60)

HEAD_FONT = 'Cambria'
BODY_FONT = 'Calibri'


# ------------------------------------------------------------------ helpers
def shade(cell, hex_fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def run(paragraph, text, *, size=11, bold=False, italic=False, color=INK,
        font=BODY_FONT, caps=False, spacing=None):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        el = OxmlElement('w:spacing')
        el.set(qn('w:val'), str(int(spacing * 20)))
        rPr.append(el)
    return r


def para(doc, *, align=None, before=0, after=0, line=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line:
        p.paragraph_format.line_spacing = line
    return p


def rule(doc, before=4, after=8):
    p = para(doc, before=before, after=after)
    pPr = p._element.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'C9A24B')
    borders.append(bottom)
    pPr.append(borders)
    return p


def cell_text(cell, text, *, size=10.5, bold=False, italic=False, color=INK,
              align=None, font=BODY_FONT):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size=size, bold=bold, italic=italic, color=color, font=font)


# ------------------------------------------------------------------- docx
def build_docx(path):
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)      # A4
    for m in ('top_margin', 'bottom_margin'):
        setattr(sec, m, Cm(1.6))
    sec.left_margin = sec.right_margin = Cm(1.8)

    base = doc.styles['Normal']
    base.font.name = BODY_FONT
    base.font.size = Pt(11)
    base.font.color.rgb = INK

    # ---- masthead
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    run(p, 'GRAPHIC ERA SCHOOL OF MANAGEMENT', size=9, bold=True,
        color=COPPER, font=HEAD_FONT, spacing=1.6)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    run(p, 'SWARAGINI, THE CULTURAL SOCIETY OF GRAPHIC ERA UNIVERSITY',
        size=8, color=GREY, font=HEAD_FONT, spacing=1.2)

    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    run(p, EVENT['name'], size=26, bold=True, color=MAROON, font=HEAD_FONT)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    run(p, EVENT['theme'], size=11.5, italic=True, color=CRIMSON, font=HEAD_FONT)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    run(p, 'PROGRAMME  \u00b7  RUNNING ORDER', size=10, bold=True,
        color=COPPER, font=HEAD_FONT, spacing=2.4)

    rule(doc, before=0, after=8)

    for line, italic in (
        (f"{EVENT['date']}  \u00b7  {EVENT['time']}", False),
        (f"{EVENT['venue']}, {EVENT['venue_sub']}", False),
        (EVENT['programme'], True),
    ):
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        run(p, line, size=10.5, italic=italic, color=INK if not italic else GREY)

    # ---- running order
    p = para(doc, before=14, after=6)
    run(p, 'RUNNING ORDER', size=11, bold=True, color=MAROON,
        font=HEAD_FONT, spacing=1.6)

    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.1), Cm(2.4), Cm(5.4), Cm(5.2), Cm(3.3)]
    heads = ['#', 'Time', 'Item', 'Details', 'Presented by']
    for i, (h, w) in enumerate(zip(heads, widths)):
        c = t.rows[0].cells[i]
        cell_text(c, h, size=9.5, bold=True, color=MAROON, font=HEAD_FONT,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else None)
        shade(c, 'F2E6CC')
        c.width = w

    for n, (time, item, detail, by) in enumerate(RUNNING_ORDER, start=1):
        cells = t.add_row().cells
        cell_text(cells[0], str(n), size=10, bold=True, color=COPPER,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[1], time, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=INK if time != TBC else GREY)
        cell_text(cells[2], item, size=10.5, bold=True)
        cell_text(cells[3], detail, size=9.5, italic=detail.startswith('To confirm'),
                  color=GREY if detail.startswith('To confirm') else INK)
        cell_text(cells[4], by, size=9.5, color=GREY if by == TBC else INK)
        for c, w in zip(cells, widths):
            c.width = w

    # ---- guests
    p = para(doc, before=14, after=6)
    run(p, 'GUESTS', size=11, bold=True, color=MAROON, font=HEAD_FONT, spacing=1.6)

    g = doc.add_table(rows=0, cols=2)
    g.style = 'Table Grid'
    for name, role in GUESTS:
        cells = g.add_row().cells
        cell_text(cells[0], name, size=10.5, bold=True)
        cell_text(cells[1], role, size=9.5, color=GREY)
        cells[0].width, cells[1].width = Cm(8.0), Cm(9.4)

    # ---- awards
    p = para(doc, before=14, after=2)
    run(p, 'AWARDS', size=11, bold=True, color=MAROON, font=HEAD_FONT, spacing=1.6)
    p = para(doc, after=6)
    run(p, AWARDS_BY, size=10, italic=True, color=CRIMSON)

    a = doc.add_table(rows=0, cols=2)
    a.style = 'Table Grid'
    for i, award in enumerate(AWARDS, start=1):
        cells = a.add_row().cells
        cell_text(cells[0], str(i), size=10, bold=True, color=COPPER,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[1], award, size=10.5)
        cells[0].width, cells[1].width = Cm(1.1), Cm(16.3)

    # ---- open questions
    p = para(doc, before=16, after=4)
    run(p, 'TO CONFIRM', size=11, bold=True, color=MAROON, font=HEAD_FONT,
        spacing=1.6)
    for q in QUERIES:
        p = para(doc, after=3)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        run(p, '\u2022  ', size=10, color=COPPER)
        # tiny bit of inline bold support for the one emphasised word
        for i, chunk in enumerate(q.replace('<b>', '\x00').replace('</b>', '\x00')
                                  .split('\x00')):
            run(p, chunk, size=10, bold=(i % 2 == 1), color=INK)

    rule(doc, before=12, after=4)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, 'Nothing on this sheet has been invented: blank times and "to confirm" '
           'entries are the points the notes did not cover.',
        size=8.5, italic=True, color=GREY)

    doc.save(path)
    return path


# --------------------------------------------------------------------- md
def build_md(path):
    L = []
    A = L.append
    A(f"# {EVENT['name']} \u2014 Programme / Running Order\n")
    A(f"*{EVENT['theme']}*\n")
    A(f"**{EVENT['date']}** &middot; {EVENT['time']}  ")
    A(f"{EVENT['venue']}, {EVENT['venue_sub']}  ")
    A(f"{EVENT['programme']}  ")
    A(f"{EVENT['hosts']}\n")
    A('## Running order\n')
    A('| # | Time | Item | Details | Presented by |')
    A('|---|---|---|---|---|')
    for n, (time, item, detail, by) in enumerate(RUNNING_ORDER, start=1):
        A(f'| {n} | {time} | **{item}** | {detail} | {by} |')
    A('\n## Guests\n')
    A('| Name | Role |')
    A('|---|---|')
    for name, role in GUESTS:
        A(f'| **{name}** | {role} |')
    A(f'\n## Awards\n\n*{AWARDS_BY}*\n')
    for i, award in enumerate(AWARDS, start=1):
        A(f'{i}. {award}')
    A('\n## To confirm\n')
    for q in QUERIES:
        A(f"- {q.replace('<b>', '**').replace('</b>', '**')}")
    A('\nNothing here has been invented: blank times and "to confirm" entries are '
      'the points the notes did not cover.\n')
    with open(path, 'w', encoding='utf-8') as fh:
        fh.write('\n'.join(L))
    return path



# --------------------------------------------------------------------- html
def esc(t):
    return (t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))


def build_html(path):
    """A designed A4 sheet, in the same palette as the invitation suite, for
    circulating and printing. The Word file remains the editable master."""
    rows = []
    for n, (time, item, detail, by) in enumerate(RUNNING_ORDER, start=1):
        show = detail and not detail.startswith('To confirm')
        rows.append(
            f'<tr><td class="n">{n}</td>'
            f'<td class="t{" tbc" if time == TBC else ""}">{esc(time)}</td>'
            f'<td class="it">{esc(item)}'
            + (f'<span class="d">{esc(detail)}</span>' if show else '')
            + f'</td><td class="by{" tbc" if by == TBC else ""}">{esc(by)}</td></tr>')

    guests = ''.join(
        f'<li><b>{esc(nm)}</b><span>{esc(role)}</span></li>' for nm, role in GUESTS)
    awards = ''.join(f'<li><i>{i}</i>{esc(a)}</li>' for i, a in enumerate(AWARDS, 1))
    n_open = len(QUERIES)

    html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<title>{esc(EVENT['name'])} &mdash; Programme</title>
<link rel="stylesheet" href="../garh-kauthig-2026-invitations/src/fonts.css">
<style>
:root {{
  --maroon:#6f1119; --crimson:#c0272d; --deep:#9c1c22; --gold:#ab8129;
  --brass:#c9a24b; --copper:#a35a1d; --ink:#33241a; --grey:#6b6257;
  --serif:'Cormorant Garamond',Georgia,serif; --caps:'Marcellus SC',serif;
  --disp:'Cinzel',serif; --sans:'Montserrat',Arial,sans-serif;
}}
*{{box-sizing:border-box}}
@page{{size:210mm 297mm;margin:0}}
html,body{{margin:0;background:#6b6257}}
.sheet{{position:relative;width:210mm;height:297mm;overflow:hidden;
  background:linear-gradient(178deg,#fdf8ec,#f9f0da 60%,#efe2c4);
  font-family:var(--serif);color:var(--ink);padding:12mm 15mm 10mm}}
.frame{{position:absolute;inset:7mm;border:0.4mm solid var(--gold);
  box-shadow:inset 0 0 0 0.1mm #fffcf0}}
.frame2{{position:absolute;inset:9mm;border:0.14mm solid var(--brass)}}
.inner{{position:relative;height:100%;display:flex;flex-direction:column}}
.host{{text-align:center;font-family:var(--sans);font-weight:700;font-size:7pt;
  letter-spacing:.22em;color:var(--copper)}}
.host2{{text-align:center;font-family:var(--sans);font-weight:500;font-size:6.4pt;
  letter-spacing:.14em;color:var(--grey);margin-top:.8mm}}
h1{{font-family:'Yellowtail',cursive;font-size:35pt;color:var(--deep);
  text-align:center;margin:3mm 0 0;line-height:1}}
.theme{{text-align:center;font-family:var(--caps);font-size:9.2pt;
  letter-spacing:.16em;color:var(--deep);margin-top:1.5mm}}
.kicker{{text-align:center;font-family:var(--disp);font-weight:700;font-size:8pt;
  letter-spacing:.3em;color:var(--gold);margin-top:2.6mm}}
.rule{{height:.16mm;background:var(--brass);margin:2.4mm 0}}
.meta{{text-align:center;font-size:9.5pt;line-height:1.5;color:var(--ink)}}
.meta em{{font-style:italic;color:var(--grey)}}
h2{{font-family:var(--disp);font-weight:700;font-size:8.4pt;letter-spacing:.24em;
  color:var(--maroon);margin:4mm 0 1.8mm;border-bottom:.16mm solid var(--brass);
  padding-bottom:1.2mm}}
table{{width:100%;border-collapse:collapse}}
td{{padding:0.88mm 1.6mm;vertical-align:top;border-bottom:.1mm solid rgba(201,162,75,.45)}}
td.n{{width:7mm;text-align:center;font-family:var(--disp);font-weight:600;
  font-size:8pt;color:var(--copper)}}
td.t{{width:20mm;font-family:var(--disp);font-size:8.2pt;white-space:nowrap}}
td.it{{font-size:10.6pt;font-weight:600}}
td.it .d{{display:block;font-weight:400;font-size:8.4pt;font-style:italic;
  color:var(--grey);margin-top:.2mm}}
td.by{{width:34mm;font-size:8.8pt;font-style:italic}}
.tbc{{color:#9a9086}}
.cols{{display:flex;gap:9mm}}
.cols>div{{flex:1}}
ul{{list-style:none;margin:0;padding:0}}
li{{font-size:9.6pt;padding:.85mm 0;border-bottom:.1mm solid rgba(201,162,75,.4);
  display:flex;gap:2mm;align-items:baseline}}
li span{{margin-left:auto;font-size:8.2pt;font-style:italic;color:var(--grey)}}
li i{{font-family:var(--disp);font-style:normal;font-weight:600;font-size:8pt;
  color:var(--copper);min-width:4mm}}
.by-line{{font-size:8.8pt;font-style:italic;color:var(--deep);margin:0 0 1.5mm}}
.onday{{margin-top:4mm}}
.onday .lines{{margin-top:2mm}}
.onday .lines div{{height:.12mm;background:rgba(201,162,75,.55);margin:5.4mm 0}}
.notes{{margin-top:auto;padding-top:2.6mm;border-top:.16mm solid var(--brass)}}
.notes h3{{font-family:var(--disp);font-weight:700;font-size:7.4pt;
  letter-spacing:.22em;color:var(--maroon);margin:0 0 1.5mm}}
.notes ul li{{border:0;font-size:7.8pt;line-height:1.35;padding:.5mm 0;
  display:list-item;list-style:none;text-indent:-2.6mm;padding-left:2.6mm}}
.notes ul li::before{{content:'\2022';color:var(--copper);padding-right:1.4mm}}
.foot{{text-align:center;font-size:7.4pt;font-style:italic;color:var(--grey);
  margin-top:2mm}}
</style></head><body>
<div class="sheet"><div class="frame"></div><div class="frame2"></div>
<div class="inner">
  <div class="host">GRAPHIC ERA SCHOOL OF MANAGEMENT</div>
  <div class="host2">SWARAGINI, THE CULTURAL SOCIETY OF GRAPHIC ERA UNIVERSITY</div>
  <h1>{esc(EVENT['name'].replace(' 2026', ''))} 2026</h1>
  <div class="theme">{esc(EVENT['theme'])}</div>
  <div class="kicker">PROGRAMME &middot; RUNNING ORDER</div>
  <div class="rule"></div>
  <div class="meta">{esc(EVENT['date'])} &nbsp;&middot;&nbsp; {esc(EVENT['time'])}<br>
    {esc(EVENT['venue'])}, {esc(EVENT['venue_sub'])}<br>
    <em>{esc(EVENT['programme'])}</em></div>

  <h2>RUNNING ORDER</h2>
  <table>{''.join(rows)}</table>

  <div class="cols">
    <div><h2>GUESTS</h2><ul>{guests}</ul></div>
    <div><h2>AWARDS</h2><p class="by-line">{esc(AWARDS_BY)}</p><ul>{awards}</ul></div>
  </div>

  <div class="onday"><h2>NOTES ON THE DAY</h2>
    <div class="lines"><div></div><div></div></div>
  </div>

  <div class="notes">
    <div class="foot">Blank times and &ldquo;to confirm&rdquo; entries mark the
      {n_open} points the organisers&rsquo; notes did not cover; nothing has been
      assumed. The full list is in the accompanying Word file.</div>
  </div>
</div></div>
<script>document.fonts.ready.then(()=>{{document.body.dataset.ready='1';}});</script>
</body></html>"""
    with open(path, 'w', encoding='utf-8') as fh:
        fh.write(html)
    return path


if __name__ == '__main__':
    d = build_docx(os.path.join(HERE, 'garh-kauthig-2026-itinerary.docx'))
    m = build_md(os.path.join(HERE, 'garh-kauthig-2026-itinerary.md'))
    h = build_html(os.path.join(HERE, 'garh-kauthig-2026-itinerary.html'))
    for f in (d, m, h):
        print(f'wrote {os.path.basename(f)}  ({os.path.getsize(f) / 1024:.0f} KB)')
    print(f'{len(RUNNING_ORDER)} programme items, {len(GUESTS)} guests, '
          f'{len(AWARDS)} awards, {len(QUERIES)} open questions')
