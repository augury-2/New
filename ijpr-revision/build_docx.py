#!/usr/bin/env python3
"""Convert manuscript.md -> submission-ready DOCX (IJPR-style formatting)."""
import re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1] if len(sys.argv) > 1 else 'manuscript.md'
DST = sys.argv[2] if len(sys.argv) > 2 else 'IJPR_Revised_Manuscript.docx'

# ---------- parse ----------
BLOCKS = []          # (kind, payload)
lines = open(SRC).read().split('\n')
i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    if not s:
        i += 1; continue
    if s.startswith('@@TITLE '):
        BLOCKS.append(('title', s[8:].strip()))
    elif s.startswith('@@FIG '):
        img, cap = s[6:].split('|', 1)
        BLOCKS.append(('fig', (img.strip(), cap.strip())))
    elif s.startswith('@@TABLE'):
        cap = s[7:].strip()
        rows = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('@@ENDTABLE'):
            if lines[i].strip():
                rows.append([c.strip() for c in lines[i].split('|')])
            i += 1
        BLOCKS.append(('table', (cap, rows)))
    elif s.startswith('#'):
        m = re.match(r'^(#+)\s*(.*)$', s)
        BLOCKS.append(('h%d' % len(m.group(1)), m.group(2).strip()))
    elif s.startswith('> '):
        BLOCKS.append(('quote', s[2:].strip()))
    elif s == '---':
        BLOCKS.append(('rule', ''))
    elif s.startswith('|'):
        # bare markdown pipe table
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells if c):
                rows.append(cells)
            i += 1
        i -= 1
        if rows:
            BLOCKS.append(('table', ('', rows)))
    elif s.startswith('- ') or s.startswith('* '):
        BLOCKS.append(('bullet', re.sub(r'^\[[ x]\]\s*', '', s[2:].strip())))
    else:
        BLOCKS.append(('p', s))
    i += 1

# ---------- document setup ----------
doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = st.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(1)

for name, sz, bold, italic in (('Heading 1', 13, True, False),
                               ('Heading 2', 12, True, False),
                               ('Heading 3', 12, False, True)):
    hs = doc.styles[name]
    hs.font.name = 'Times New Roman'; hs.font.size = Pt(sz)
    hs.font.bold = bold; hs.font.italic = italic
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.paragraph_format.line_spacing = 2.0
    hs.paragraph_format.keep_with_next = True

INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')

def add_runs(par, text):
    """Render **bold** and *italic* inline markers."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)

def shade(cell, hexcolor):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor); cell._tc.get_or_add_tcPr().append(el)

def add_table(cap, rows):
    if cap:
        c = doc.add_paragraph(); c.paragraph_format.space_before = Pt(12)
        c.paragraph_format.line_spacing = 1.0; c.paragraph_format.keep_with_next = True
        r = c.add_run(cap); r.bold = True; r.font.size = Pt(10)
    ncol = max(len(x) for x in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ''
            p = cells[ci].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_runs(p, txt)
            for run in p.runs:
                run.font.size = Pt(9); run.font.name = 'Times New Roman'
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                shade(cells[ci], 'EFEFEF')
    doc.add_paragraph().paragraph_format.line_spacing = 1.0

# ---------- render ----------
for kind, payload in BLOCKS:
    if kind == 'title':
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(payload); r.bold = True; r.font.size = Pt(15)
    elif kind == 'h1':
        doc.add_paragraph(payload, style='Heading 1')
    elif kind == 'h2':
        doc.add_paragraph(payload, style='Heading 2')
    elif kind == 'h3':
        doc.add_paragraph(payload, style='Heading 3')
    elif kind == 'quote':
        p = doc.add_paragraph(); pfm = p.paragraph_format
        pfm.left_indent = Inches(0.5); pfm.right_indent = Inches(0.5)
        pfm.space_before = Pt(6); pfm.space_after = Pt(6); pfm.line_spacing = 1.5
        r = p.add_run(payload); r.italic = True
    elif kind == 'bullet':
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 2.0
        add_runs(p, payload)
    elif kind == 'fig':
        img, cap = payload
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        try:
            p.add_run().add_picture(img, width=Inches(5.6))
        except Exception:
            p.add_run('[%s not found]' % img).italic = True
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.line_spacing = 1.0; c.paragraph_format.space_after = Pt(12)
        cr = c.add_run(cap); cr.font.size = Pt(10)
    elif kind == 'table':
        add_table(*payload)
    elif kind == 'rule':
        p = doc.add_paragraph()
        pbdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '999999')
        pbdr.append(bot); p._p.get_or_add_pPr().append(pbdr)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.0)
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, payload)

doc.save(DST)
words = sum(len(t.split()) for k, t in BLOCKS if k in ('p', 'quote', 'bullet', 'title'))
print('wrote %s | blocks=%d | prose words=%d' % (DST, len(BLOCKS), words))
