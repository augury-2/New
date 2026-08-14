#!/usr/bin/env python3
"""
Build a DOCX with genuine Word tracked changes (w:ins / w:del) by diffing the
ORIGINAL manuscript's paragraph text against the REVISED manuscript's paragraph text.

Paragraph-level alignment via difflib on normalised text, then word-level diff
inside aligned pairs. Output opens in Word with Review > All Markup showing
every insertion and deletion.
"""
import re, sys, difflib
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table

ORIG = 'New_original.docx'
REV_MD = 'manuscript.md'
DST = 'IJPR_Revised_Manuscript_TrackedChanges.docx'
AUTHOR = 'Kiro (editorial revision)'
STAMP = '2026-08-14T00:00:00Z'

# ---------------- read original paragraphs ----------------
od = Document(ORIG)
orig_paras = []
for child in od.element.body.iterchildren():
    if child.tag.endswith('}p'):
        t = Paragraph(child, od).text.strip()
        if t:
            orig_paras.append(t)

# ---------------- read revised blocks ----------------
BLOCKS = []
lines = open(REV_MD).read().split('\n')
i = 0
while i < len(lines):
    s = lines[i].strip()
    if not s:
        i += 1; continue
    if s.startswith('@@TITLE '):
        BLOCKS.append(('title', s[8:].strip()))
    elif s.startswith('@@FIG '):
        img, cap = s[6:].split('|', 1)
        BLOCKS.append(('fig', (img.strip(), cap.strip())))
    elif s.startswith('@@TABLE'):
        cap = s[7:].strip(); rows = []; i += 1
        while i < len(lines) and not lines[i].strip().startswith('@@ENDTABLE'):
            if lines[i].strip():
                rows.append([c.strip() for c in lines[i].split('|')])
            i += 1
        BLOCKS.append(('table', (cap, rows)))
    elif s.startswith('#'):
        m = re.match(r'^(#+)\s*(.*)$', s)
        BLOCKS.append(('h%d' % len(m.group(1)), m.group(2).strip()))
    elif s.startswith('> '):
        BLOCKS.append(('quote', s[2:].strip()))
    elif s.startswith('- '):
        BLOCKS.append(('bullet', s[2:].strip()))
    else:
        BLOCKS.append(('p', s))
    i += 1

TEXTKINDS = ('p', 'quote', 'bullet', 'title', 'h1', 'h2', 'h3')
rev_idx = [k for k, b in enumerate(BLOCKS) if BLOCKS[k][0] in TEXTKINDS]
rev_paras = [BLOCKS[k][1] for k in rev_idx]

def strip_md(t):
    return re.sub(r'\*+', '', t)

def norm(t):
    return re.sub(r'[^a-z0-9 ]', '', strip_md(t).lower())

# ---------------- paragraph alignment ----------------
sm = difflib.SequenceMatcher(None, [norm(x) for x in orig_paras],
                             [norm(x) for x in rev_paras], autojunk=False)
# map revised-para position -> matched original text (or None = wholly new)
match_for = {}
consumed = set()
for a, b, n in sm.get_matching_blocks():
    for k in range(n):
        match_for[b + k] = a + k
        consumed.add(a + k)
# fuzzy-match remaining revised paragraphs to unconsumed originals
unmatched_orig = [j for j in range(len(orig_paras)) if j not in consumed]
for bi in range(len(rev_paras)):
    if bi in match_for:
        continue
    best, bestscore = None, 0.0
    nb = norm(rev_paras[bi])
    if len(nb) < 40:
        continue
    for j in unmatched_orig:
        na = norm(orig_paras[j])
        if abs(len(na) - len(nb)) > max(len(na), len(nb)) * 0.75:
            continue
        r = difflib.SequenceMatcher(None, na, nb).quick_ratio()
        if r > bestscore:
            bestscore, best = r, j
    if best is not None and bestscore >= 0.55:
        real = difflib.SequenceMatcher(None, norm(orig_paras[best]), nb).ratio()
        if real >= 0.45:
            match_for[bi] = best
            unmatched_orig.remove(best)

deleted_orig = [j for j in range(len(orig_paras)) if j not in set(match_for.values())]

# ---------------- document scaffolding ----------------
doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(0)
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Inches(1)
for name, sz, bold, italic in (('Heading 1', 13, True, False),
                               ('Heading 2', 12, True, False),
                               ('Heading 3', 12, False, True)):
    hs = doc.styles[name]
    hs.font.name = 'Times New Roman'; hs.font.size = Pt(sz)
    hs.font.bold = bold; hs.font.italic = italic
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12); hs.paragraph_format.space_after = Pt(6)

RID = [1000]
def rid():
    RID[0] += 1
    return str(RID[0])

def mkrun(text, bold=False, italic=False, size=None, strike=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), 'Times New Roman'); f.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(f)
    if bold:   rPr.append(OxmlElement('w:b'))
    if italic: rPr.append(OxmlElement('w:i'))
    if strike: rPr.append(OxmlElement('w:strike'))
    if size:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    r.append(rPr)
    return r, rPr, text

def append_plain(par, text, bold=False, italic=False, size=None):
    r, rPr, _ = mkrun(text, bold, italic, size)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); par._p.append(r)

def append_ins(par, text, bold=False, italic=False, size=None):
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), rid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), STAMP)
    r, rPr, _ = mkrun(text, bold, italic, size)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); ins.append(r); par._p.append(ins)

def append_del(par, text, size=None):
    d = OxmlElement('w:del')
    d.set(qn('w:id'), rid()); d.set(qn('w:author'), AUTHOR); d.set(qn('w:date'), STAMP)
    r, rPr, _ = mkrun(text, size=size)
    t = OxmlElement('w:delText'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); d.append(r); d_par = par._p; d_par.append(d)

def mark_para_inserted(par):
    """Mark the paragraph mark itself as an insertion."""
    pPr = par._p.get_or_add_pPr()
    rPr = OxmlElement('w:rPr')
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), rid()); ins.set(qn('w:author'), AUTHOR); ins.set(qn('w:date'), STAMP)
    rPr.append(ins); pPr.append(rPr)

def mark_para_deleted(par):
    pPr = par._p.get_or_add_pPr()
    rPr = OxmlElement('w:rPr')
    d = OxmlElement('w:del')
    d.set(qn('w:id'), rid()); d.set(qn('w:author'), AUTHOR); d.set(qn('w:date'), STAMP)
    rPr.append(d); pPr.append(rPr)

WORDS = re.compile(r'\S+\s*')

def emit_worddiff(par, old, new, bold=False, italic=False, size=None):
    """Word-level tracked diff of old -> new inside one paragraph."""
    a = WORDS.findall(old); b = WORDS.findall(new)
    s = difflib.SequenceMatcher(None, [w.strip().lower() for w in a],
                               [w.strip().lower() for w in b], autojunk=False)
    for tag, i1, i2, j1, j2 in s.get_opcodes():
        if tag == 'equal':
            append_plain(par, ''.join(b[j1:j2]), bold, italic, size)
        elif tag == 'delete':
            append_del(par, ''.join(a[i1:i2]), size)
        elif tag == 'insert':
            append_ins(par, ''.join(b[j1:j2]), bold, italic, size)
        else:
            append_del(par, ''.join(a[i1:i2]), size)
            append_ins(par, ''.join(b[j1:j2]), bold, italic, size)

def shade(cell, hexcolor):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor); cell._tc.get_or_add_tcPr().append(el)

# ---------------- banner ----------------
b = doc.add_paragraph(); b.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = b.add_run('TRACKED-CHANGES COPY — open in Word and choose Review > All Markup. '
              'Insertions and deletions are recorded as Word revisions attributed to "%s". '
              'Deleted paragraphs from the original are grouped in Appendix Z at the end.' % AUTHOR)
r.italic = True; r.font.size = Pt(9)
doc.add_paragraph()

# ---------------- render revised, with tracked diff ----------------
pos_in_rev = {k: n for n, k in enumerate(rev_idx)}
stats = {'ins_paras': 0, 'edited_paras': 0, 'kept_paras': 0}

for bidx, (kind, payload) in enumerate(BLOCKS):
    if kind == 'table':
        cap, rows = payload
        c = doc.add_paragraph(); c.paragraph_format.line_spacing = 1.0
        c.paragraph_format.space_before = Pt(12)
        append_ins(c, cap, bold=True, size=10)
        mark_para_inserted(c)
        ncol = max(len(x) for x in rows)
        t = doc.add_table(rows=0, cols=ncol)
        t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(ncol):
                txt = row[ci] if ci < len(row) else ''
                p = cells[ci].paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                if txt:
                    append_ins(p, txt, bold=(ri == 0), size=9)
                if ri == 0:
                    shade(cells[ci], 'EFEFEF')
        doc.add_paragraph().paragraph_format.line_spacing = 1.0
        continue
    if kind == 'fig':
        img, cap = payload
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(img, width=Inches(5.2))
        except Exception:
            append_plain(p, '[%s]' % img, italic=True)
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.line_spacing = 1.0
        append_ins(c, cap, size=10)
        continue

    text = payload
    style = {'h1': 'Heading 1', 'h2': 'Heading 2', 'h3': 'Heading 3',
             'bullet': 'List Bullet'}.get(kind)
    par = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if kind == 'title':
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if kind == 'quote':
        par.paragraph_format.left_indent = Inches(0.5)
        par.paragraph_format.right_indent = Inches(0.5)
    par.paragraph_format.space_after = Pt(6)

    bold = kind == 'title'
    size = 15 if kind == 'title' else None
    italic = kind == 'quote'
    plain = strip_md(text)
    n = pos_in_rev[bidx]
    if n in match_for:
        old = orig_paras[match_for[n]]
        if norm(old) == norm(plain):
            append_plain(par, plain, bold, italic, size)
            stats['kept_paras'] += 1
        else:
            emit_worddiff(par, old, plain, bold, italic, size)
            stats['edited_paras'] += 1
    else:
        append_ins(par, plain, bold, italic, size)
        mark_para_inserted(par)
        stats['ins_paras'] += 1

# ---------------- deleted originals ----------------
if deleted_orig:
    doc.add_page_break()
    h = doc.add_paragraph(style='Heading 1')
    append_ins(h, 'Appendix Z. Original passages removed or wholly rewritten'); mark_para_inserted(h)
    n = doc.add_paragraph()
    append_ins(n, 'The passages below appeared in the original manuscript and are marked as '
                  'deletions. Each was either replaced by repositioned text elsewhere in this '
                  'document or removed as redundant. Accepting all changes removes this appendix.',
               italic=True, size=10)
    mark_para_inserted(n)
    for j in deleted_orig:
        p = doc.add_paragraph()
        append_del(p, orig_paras[j])
        mark_para_deleted(p)

doc.save(DST)
print('wrote %s' % DST)
print('  revised paragraphs: %d  (new: %d, edited-with-diff: %d, unchanged: %d)'
      % (len(rev_paras), stats['ins_paras'], stats['edited_paras'], stats['kept_paras']))
print('  original paragraphs: %d  (deleted/rewritten: %d)' % (len(orig_paras), len(deleted_orig)))
